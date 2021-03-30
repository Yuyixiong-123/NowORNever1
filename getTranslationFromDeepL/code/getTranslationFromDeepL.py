# -*- coding: utf-8 -*-
"""
Created on Thu Mar 25 10:32:33 2021

@author: YU Yixiong
"""

import docx
import openpyxl
from selenium import webdriver
import time
from webdriver_manager.chrome import ChromeDriverManager
import pyperclip
import Config
import tqdm
from selenium.webdriver.chrome.options import Options

def ChromeDriverNOBrowser():
   chrome_options = Options()
   chrome_options.add_argument('--headless')
   chrome_options.add_argument('--disable-gpu')
   driverChrome = webdriver.Chrome(chrome_options=chrome_options)
   return driverChrome

def getTranFromDeepL(engPara,sleepTime):
    browser=webdriver.Chrome()
    browser.get(r'https://www.deepl.com/translator#en/zh/'+engPara)
    
    time.sleep(sleepTime) 
    browser.find_element_by_class_name("lmt__target_toolbar__copy").click()
    time.sleep(0.5) 
    
    # print(0)
    copyStr=pyperclip.paste()
    print(copyStr)
    # print(1)
    # time.sleep(1) 
    browser.close()
    return copyStr

def isTitle(result):
    while result[0]==' ' or ord(result[0])==160:
          result=result[1:]
    if ord(result[0])==52:
        return True
    else:
        return False
    
def getHeadingLevel(result):
    headSeri=''
    counter=-1 
    
    for s in range(len(result)):
         if result[s]==' ':
              headSeri=result[:s]
              break
    for s in headSeri:
         if s==".":
              counter+=1
     # print(headSeri)
    return counter
     
def modifyTheTitle(titleStr):
    prefix=''
    suffix=''
    while(titleStr[0]==' ' or ord(titleStr[0])==160):
          titleStr=titleStr[1:]
    for s in range(len(titleStr)):
        if titleStr[s]==' ' or ord(titleStr[s])==160:
             prefix=titleStr[:s]
             print(prefix)
             suffix=titleStr[s:]
             break
    while(suffix[0]==' ' or ord(suffix[0])==160):
          suffix=suffix[1:]
          # print(suffix)
    return prefix+' '+suffix

def getSleepTime(string):
    l=len(string)
    if l<50:
        return 6
    if l>700:
        return 10
    else:
        return 8
if __name__ =="__main__":
     
     document=docx.Document(Config.templatePath)
     wb=openpyxl.load_workbook(Config.paraElementsPath)
     sheet=wb["Sheet1"]
     
     for i in range(1,sheet.max_row+1):#max row equal 68
     # for i in range(1,50):
          print(i)
     # for i in range(1,sheet.max_row+1):
          engPara=str(sheet.cell(i,1).value)
          try:
              result=getTranFromDeepL(engPara,getSleepTime(engPara))
              headLevel=getHeadingLevel(engPara)
              sheet.cell(i,2).value=result
              
              if isTitle(engPara) ==True :
                   engPara=modifyTheTitle(engPara)
                   headLevel=getHeadingLevel(engPara)
                   document.add_heading(engPara,headLevel)
                   document.add_paragraph(result)
              else:
                   document.add_paragraph('\n    '+engPara)
                   document.add_paragraph(result)
          except:
              document.add_paragraph('\n    '+engPara)
              pass
     wb.save(Config.paraElementsPath)
     document.save(Config.savePath)
