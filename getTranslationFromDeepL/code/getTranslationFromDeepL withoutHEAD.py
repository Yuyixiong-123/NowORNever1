# -*- coding: utf-8 -*-
"""
Created on Thu Mar 25 10:32:33 2021

@author: YU Yixiong
"""

import docx
import openpyxl
from selenium import webdriver
import time
import Config
from selenium.webdriver.chrome.options import Options
import re

def getTranFromDeepL(engPara,sleepTime):
    options = Options()
    options.add_argument('--headless')
    
    browser = webdriver.Chrome(options=options) 
    # engPara="Not the answer you're looking for? Browse other questions tagged python unicode encode or ask your own question. "
    browser.get(r'https://www.deepl.com/translator#en/zh/'+engPara)
    # print('sleep')
    time.sleep(sleepTime)
    # print('wake up')
    fs=(browser.page_source)
    regex=re.compile('lmt__translations_as_text__text_btn">(.*?)</button><button class="lmt')
    x=regex.findall(fs)
    # print(x)
    browser.close()
    return x[0]

def isTitle(result):
    while result[0]==' ' or ord(result[0])==160:
          result=result[1:]
    if ord(result[0])==53:#已经改成了为5的时候
        return True
    else:
        return False
    
def getHeadingLevel(result):
    headSeri=''
    counter=-1 
    
    for s in range(len(result)):
         if result[s]==' ':
              headSeri=result[:s]
              break
    for s in headSeri:
         if s==".":
              counter+=1
     # print(headSeri)
    return counter
     
def modifyTheTitle(titleStr):
    prefix=''
    suffix=''
    while(titleStr[0]==' ' or ord(titleStr[0])==160):
          titleStr=titleStr[1:]
    for s in range(len(titleStr)):
        if titleStr[s]==' ' or ord(titleStr[s])==160:
             prefix=titleStr[:s]
             print(prefix)
             suffix=titleStr[s:]
             break
    while(suffix[0]==' ' or ord(suffix[0])==160):
          suffix=suffix[1:]
          # print(suffix)
    return prefix+' '+suffix

def getSleepTime(string):
    l=len(string)
    if l<50:
        return 12
    if l>700:
        return 20
    else:
        return 16

def saveData(wb,document):
    wb.save(Config.paraElementsPath)
    document.save(Config.savePath)
    
    document=docx.Document(Config.savePath)
    
    wb=openpyxl.load_workbook(Config.paraElementsPath)
    sheet=wb["Sheet1"]
    return sheet

if __name__ =="__main__":
     
     document=docx.Document(Config.templatePath)
     wb=openpyxl.load_workbook(Config.paraElementsPath)
     sheet=wb["Sheet1"]
     
     # for i in range(250,sheet.max_row+1):#max row equal 68
     for i in range(1,sheet.max_row+1):
          print(i)
     # for i in range(1,sheet.max_row+1):
          engPara=str(sheet.cell(i,1).value)
          try:
              result=getTranFromDeepL(engPara,getSleepTime(engPara))
              headLevel=getHeadingLevel(engPara)
              sheet.cell(i,2).value=result
              
              if isTitle(engPara) ==True :
                   engPara=modifyTheTitle(engPara)
                   headLevel=getHeadingLevel(engPara)
                   document.add_heading(engPara,headLevel)
                   document.add_paragraph(result)
                   sheet=saveData(wb, document)
              else:
                   document.add_paragraph('\n    '+engPara)
                   document.add_paragraph(result)
                   sheet=saveData(wb, document)
          except:
              document.add_paragraph('\n    '+engPara)
              pass
     wb.save(Config.paraElementsPath)
     document.save(Config.savePath)
