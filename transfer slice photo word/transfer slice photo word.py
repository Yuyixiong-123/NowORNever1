# -*- coding: utf-8 -*-
"""
Created on Thu Nov 26 10:47:33 2020
No pain no gain
@author: yu yixiong
"""


name="习近平党史观"
path=r"C:/Users/YU Yixiong/Desktop/习近平党史观"

from aip import AipOcr
import os
from PIL import Image
from docx import Document
from docx.shared import Inches

APP_ID='老老实实抱着我'
API_KEY='517d60b47eed4ebfa794dcf667f5eec3'
SECRET_KEY='e054e6b427d0468296e3436893c5e25b'
client=AipOcr(APP_ID, API_KEY, SECRET_KEY)
 
def get_file_content(filePath):
    with open(filePath, 'rb') as fp:
        return fp.read()
 
def image2text(fileName):
    image = get_file_content(fileName)
    dic_result = client.basicGeneral(image)
    res = dic_result['words_result']
    result = ''
    r=1
    for m in res:
        # print(str(m['words']))
        result = result + str(m['words'])
        if r==1:
            result+="\n"
        r+=1
    return result

doc = Document("template.docx") 
doc.save(name+'.docx')
for i in range(len(os.listdir(path))):
    doc=Document(name+'.docx')
    index=path+"/ (%i).png"%(i+1)
    getresult = image2text(index)
    
    # to delete the whatnot info that prohibit python-docx insert image
    f = Image.open(index)#你的图片文件
    f.save(index)#替换掉你的图片文件
    f.close()
    
    doc.add_picture(index,width=Inches(6))
    paragraph = doc.add_paragraph(getresult)
    doc.save(name+'.docx')
    
    print('页面'+str(i)+'已完成')



