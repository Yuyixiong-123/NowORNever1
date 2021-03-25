from docx import Document
# from docx.shared import Inches

from PIL import Image
f = Image.open('1.jpg')#你的图片文件
f.save('1.jpg')#替换掉你的图片文件
f.close()

# from docx import Document

doc = Document()  # 以默认模板建立文档对象

doc = Document('a.docx')  # 读取a.docx文档，建立文档对象

from docx.shared import Inches, Pt

from docx.oxml.ns import qn


def chg_font(obj, fontname='微软雅黑', size=None):
    ## 设置字体函数

    obj.font.name = fontname

    obj._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)

    if size and isinstance(size, Pt):
        obj.font.size = size

# def writeTxt(txt):
distance = Inches(0.3)

sec = doc.sections[0]  # sections对应文档中的“节”

sec.left_margin = distance  # 以下依次设置左、右、上、下页面边距

sec.right_margin = distance

sec.top_margin = distance

sec.bottom_margin = distance

sec.page_width = Inches(12)  # 设置页面宽度

sec.page_height = Inches(20)  # 设置页面高度

##设置默认字体

chg_font(doc.styles['Normal'], fontname='宋体')

##步骤三：
##1.添加段落文本
paragraph = doc.add_paragraph('text....')

ph_format = paragraph.paragraph_format

ph_format.space_before = Pt(10)  # 设置段前间距

ph_format.space_after = Pt(12)  # 设置段后间距

ph_format.line_spacing = Pt(19)  # 设置行间距

##2.添加表格
# run.add_picture('a.png')  # 插入图像，可以是内存中的图像，width=Inches(1.0)指定宽度。


##3.在文档中添加图像；
doc.add_picture("1.jpg",width=Inches(10.0))

##步骤四：
doc.save('a.docx')    #保存图像

